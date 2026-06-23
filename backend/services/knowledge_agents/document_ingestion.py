import os
import io
import re
from typing import List, Optional
from datetime import datetime

from .security import validate_path, validate_file_type, SecurityError
from .models import KnowledgeSource, SourceType, SourceStatus, Chunk


PDF_AVAILABLE = False
DOCX_AVAILABLE = False
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pdfplumber
        PDF_AVAILABLE = True
    except ImportError:
        pass

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    pass


class DocumentIngestion:
    def __init__(self):
        self.supported_types = {
            ".pdf": self._extract_pdf,
            ".docx": self._extract_docx,
            ".doc": self._extract_docx,
            ".txt": self._extract_text,
            ".md": self._extract_markdown,
            ".mdx": self._extract_markdown,
        }

    def ingest(self, file_path: str, source: Optional[KnowledgeSource] = None) -> dict:
        abs_path = validate_path(file_path)
        ext = validate_file_type(file_path)

        extractor = self.supported_types.get(ext)
        if not extractor:
            raise SecurityError(f"Unsupported file type: {ext}")

        text, metadata = extractor(abs_path)

        result = {
            "text": text,
            "metadata": {
                "filename": os.path.basename(abs_path),
                "file_path": abs_path,
                "file_size": os.path.getsize(abs_path),
                "file_type": ext,
                "source_type": SourceType.DOCUMENT.value,
                "extracted_at": datetime.utcnow().isoformat(),
                **metadata,
            },
            "source_type": SourceType.DOCUMENT,
        }

        if source:
            source.metadata.update(result["metadata"])

        return result

    def _extract_pdf(self, path: str) -> tuple:
        text_parts = []
        metadata = {}
        try:
            if not PDF_AVAILABLE:
                raise ImportError("No PDF library available")
            try:
                with open(path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    metadata["page_count"] = len(reader.pages)
                    metadata["title"] = reader.metadata.title if reader.metadata else None
                    for page in reader.pages:
                        page_text = page.extract_text() or ""
                        text_parts.append(page_text)
            except NameError:
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    metadata["page_count"] = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        text_parts.append(page_text)
        except Exception as e:
            text_parts.append(f"[PDF extraction error: {e}]")

        return "\n".join(text_parts), metadata

    def _extract_docx(self, path: str) -> tuple:
        text_parts = []
        metadata = {}
        if DOCX_AVAILABLE:
            try:
                doc = docx.Document(path)
                metadata["paragraph_count"] = len(doc.paragraphs)
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        text_parts.append(" | ".join(cells))
            except Exception as e:
                text_parts.append(f"[DOCX extraction error: {e}]")
        else:
            text_parts.append(f"[DOCX extraction requires python-docx]")
        return "\n".join(text_parts), metadata

    def _extract_text(self, path: str) -> tuple:
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
        except Exception as e:
            text = f"[Text extraction error: {e}]"
        return text, {}

    def _extract_markdown(self, path: str) -> tuple:
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            headers = re.findall(r"^#{1,6}\s+(.+)$", text, re.MULTILINE)
            return text, {"headers": headers, "header_count": len(headers)}
        except Exception as e:
            return f"[Markdown extraction error: {e}]", {}

    def get_supported_extensions(self) -> List[str]:
        return list(self.supported_types.keys())
